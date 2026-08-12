"""Document parsing: RTF/DOCX/TXT/PDF -> clean text (with formula OCR)."""

import re
import subprocess
import tempfile
from pathlib import Path


def extract_text(file_path: str, ocr_formulas: bool = True) -> str:
    """Extract text; optionally insert LaTeX formulas (via pix2text server)."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    text = ""
    if suffix == ".rtf":
        text = _extract_rtf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix == ".pdf":
        text = _extract_pdf(path)
        if ocr_formulas:
            text = _merge_pdf_formulas(path, text)
    elif suffix in (".txt", ".md"):
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        text = _extract_rtf(path)
    return text


def _merge_pdf_formulas(pdf_path: Path, text: str) -> str:
    """OCR formulas from PDF pages and merge LaTeX into text.

    Only processes pages that likely contain formulas (cheap heuristic from pypdf text).
    Maximum 30 pages to keep processing time reasonable (~5 min total).
    """
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(pdf_path))
        pages_text = text.split("\f")
    except Exception:
        return text

    import formula_ocr
    import logging
    _log = logging.getLogger("ingestion")

    # Heuristic: pages with formula-like content
    formula_indicators = [
        "определяется как:", "рассчитывается по", "по формуле", "по формулам",
        "\\sum", "\\max(", "\\min(", "\\frac", "правил определения",
        "max(", "min(", "/)", "×", "+", "=",
    ]
    page_candidates = []
    for i in range(min(len(pages_text), len(reader.pages))):
        pt = pages_text[i] if i < len(pages_text) else ""
        score = sum(1 for kw in formula_indicators if kw in pt)
        if score >= 2:
            page_candidates.append(i)

    MAX_PAGES = 30
    if len(page_candidates) > MAX_PAGES:
        page_candidates = page_candidates[:MAX_PAGES]

    if not page_candidates:
        return text

    _log.info("Formula OCR: %d candidate pages (from %d total)", len(page_candidates), len(reader.pages))
    parts = list(pages_text)
    for idx in page_candidates:
        page_num = idx + 1
        page_text = parts[idx]
        try:
            png_bytes = _render_page_to_png(pdf_path, page_num)
        except Exception:
            continue
        if not png_bytes:
            continue
        formulas = formula_ocr.extract_formulas(png_bytes)
        if formulas:
            latex = "\n".join(f"$${f}$$" for f in sorted(set(formulas), key=len, reverse=True))
            parts[idx] = page_text.rstrip() + "\n\n[ФОРМУЛЫ]\n" + latex + "\n"
            _log.info("  стр %d: %d формул", page_num, len(formulas))
        else:
            _log.info("  стр %d: нет формул", page_num)
    return "\f".join(parts)


def _render_page_to_png(pdf_path: Path, page_num: int) -> bytes | None:
    """Render a PDF page to PNG via pdftoppm. Returns PNG bytes."""
    with tempfile.TemporaryDirectory() as tmp:
        out_prefix = str(Path(tmp) / "page")
        res = subprocess.run(
            ["pdftoppm", "-f", str(page_num), "-l", str(page_num), "-r", "200",
             "-png", str(pdf_path), out_prefix],
            capture_output=True, timeout=120,
        )
        if res.returncode != 0:
            return None
        import glob
        files = glob.glob(out_prefix + "*.png")
        if not files:
            return None
        return Path(files[0]).read_bytes()


def _extract_rtf(path: Path) -> str:
    """Extract text from RTF. Handles cp1251 hex escapes and tables."""
    try:
        from striprtf.striprtf import rtf_to_text
        raw = path.read_text(encoding="cp1251", errors="replace")
        return rtf_to_text(raw)
    except ImportError:
        pass

    # Fallback manual stripper
    raw = path.read_bytes().decode("latin-1")
    text = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: _hex_to_char(m.group(1)), raw)
    text = re.sub(r"\\[a-zA-Z]+\d*\s?", " ", text)
    text = re.sub(r"\{[^{}]*\}", "", text)
    text = text.replace("\\par", "\n").replace("\\cell", " | ").replace("\\row", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text


def _hex_to_char(hex_str: str) -> str:
    try:
        return bytes([int(hex_str, 16)]).decode("cp1251")
    except Exception:
        return "?"


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts)
    except ImportError:
        return path.read_text(encoding="utf-8", errors="replace")


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\f".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        return ""


def chunk_text(text: str, chunk_size: int = 1500, overlap: int = 150) -> list[str]:
    """Split text into overlapping chunks, preferring article/paragraph boundaries."""
    # Normalize whitespace but keep paragraphs
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Try splitting by sections first
    sections = re.split(r"(?=(?:Статья|Пункт|Раздел|Приложение|Глава)\s*\d)", text)
    sections = [s.strip() for s in sections if s.strip()]

    if len(sections) > 1 and all(len(s) > 100 for s in sections):
        chunks: list[str] = []
        for section in sections:
            chunks.extend(_chunk_by_size(section, chunk_size, overlap))
        return chunks
    return _chunk_by_size(text, chunk_size, overlap)


def chunk_has_formula(chunk: str) -> bool:
    """True if chunk contains a LaTeX formula block."""
    return bool(re.search(r"\$\$.*?\$\$", chunk, re.DOTALL))


def _chunk_by_size(text: str, chunk_size: int, overlap: int) -> list[str]:
    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            # try to break at paragraph boundary
            boundary = text.rfind("\n", start + chunk_size // 2, end)
            if boundary == -1:
                boundary = text.rfind(" ", start + chunk_size // 2, end)
            if boundary != -1:
                end = boundary
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = max(end - overlap, start + 1)
        if start >= len(text):
            break
    return chunks


def parse_metadata_from_filename(filename: str) -> dict:
    base = filename.rsplit(".", 1)[0].replace("_", " ")
    number = re.search(r"N\s*(\d+)", base)
    date = re.search(r"от\s+(\d{2}\.\d{2}\.\d{4})", base)
    revision = re.search(r"ред\.?\s*от\s*(\d{2}\.\d{2}\.\d{4})", base)
    title = re.sub(r"\s+от\s+\d{2}\.\d{2}\.\d{4}.*", "", base)
    title = re.sub(r"\s+N\s*\d+.*", "", title).strip()
    return {
        "title": title or filename,
        "doc_number": number.group(1) if number else None,
        "doc_date": date.group(1) if date else None,
        "revision": revision.group(1) if revision else None,
    }
