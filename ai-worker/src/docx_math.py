"""Math extraction from DOCX/RTF: oMath -> LaTeX (no OCR) + WMF/EMF images -> pix2text.

Two paths for formulas in Office documents:

1. Native Word equations (`m:oMath`) are converted to LaTeX directly from
   `word/document.xml` — no OCR, no errors, fast.
2. Legacy formulas (MathType and friends) are embedded as WMF/EMF pictures
   (`w:drawing` / `w:pict` / `w:object` inside runs, or `{\\pict\\wmetafile8}`
   in RTF). These are rasterized to PNG via LibreOffice (soffice) and sent to
   the pix2text server for recognition (MFR).

If pix2text or LibreOffice is unavailable, image formulas are skipped
gracefully (same behaviour as before this module). Text extraction never fails.
"""

import logging
import re
import subprocess
import tempfile
import zipfile
from pathlib import Path

logger = logging.getLogger("docx-math")

# --- namespaces ----------------------------------------------------------

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

_FORMULA_IMAGE_EXTS = ("wmf", "emf")


def _local(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _val(el, ns: str = M_NS, attr: str = "val") -> str:
    return el.get("{%s}%s" % (ns, attr)) or ""


# --- oMath -> LaTeX -------------------------------------------------------

_LATEX_ESCAPES = {
    "\\": r"\backslash{}",
    "^": r"\^{}",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "%": r"\%",
    "&": r"\&",
    "#": r"\#",
    "$": r"\$",
    "~": r"\textasciitilde{}",
}


def _latex_escape(text: str) -> str:
    return "".join(_LATEX_ESCAPES.get(c, c) for c in text)


def omath_to_latex(el) -> str:
    """Convert an m:oMath / m:oMathPara element to LaTeX text."""
    parts = []
    for child in el:
        name = _local(child.tag)
        handler = _OMATH_HANDLERS.get(name)
        if handler is not None:
            parts.append(handler(child))
    return "".join(parts)


def _h_t(el) -> str:
    return _latex_escape(el.text or "")


def _h_r(el) -> str:
    text = "".join(_h_t(c) for c in el if _local(c.tag) == "t")
    style = None
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "rPr":
            for n2, c2 in ((_local(c2.tag), c2) for c2 in c):
                if n2 == "sty":
                    style = _val(c2)
    if style in ("i", "bi"):
        return r"\mathit{" + text + "}"
    if style == "b":
        return r"\mathbf{" + text + "}"
    return text


def _h_f(el) -> str:
    num = den = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "num":
            num = omath_to_latex(c)
        elif name == "den":
            den = omath_to_latex(c)
    return r"\frac{%s}{%s}" % (num, den)


def _h_sSup(el) -> str:
    base = sup = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "e":
            base = omath_to_latex(c)
        elif name == "sup":
            sup = omath_to_latex(c)
    return "{%s}^{%s}" % (base, sup)


def _h_sSub(el) -> str:
    base = sub = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "e":
            base = omath_to_latex(c)
        elif name == "sub":
            sub = omath_to_latex(c)
    return "{%s}_{%s}" % (base, sub)


def _h_sSubSup(el) -> str:
    base = sub = sup = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "e":
            base = omath_to_latex(c)
        elif name == "sub":
            sub = omath_to_latex(c)
        elif name == "sup":
            sup = omath_to_latex(c)
    return "{%s}_{%s}^{%s}" % (base, sub, sup)


def _h_sPre(el) -> str:
    sub = sup = base = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "sub":
            sub = omath_to_latex(c)
        elif name == "sup":
            sup = omath_to_latex(c)
        elif name == "e":
            base = omath_to_latex(c)
    return "{}_{%s}^{%s}{%s}" % (sub, sup, base)


_NARY_OPS = {
    "∑": r"\sum",
    "∏": r"\prod",
    "∐": r"\coprod",
    "∫": r"\int",
    "∮": r"\oint",
    "∪": r"\bigcup",
    "⋃": r"\bigcup",
    "∩": r"\bigcap",
    "⋂": r"\bigcap",
    "⊕": r"\bigoplus",
    "⊗": r"\bigotimes",
    "⊙": r"\bigodot",
    "∨": r"\bigvee",
    "∧": r"\bigwedge",
    "⨁": r"\bigoplus",
    "⨂": r"\bigotimes",
}

_NARY_LIMITS_OPS = {r"\sum", r"\prod", r"\coprod", r"\bigcup", r"\bigcap",
                    r"\bigoplus", r"\bigotimes", r"\bigodot", r"\bigvee", r"\bigwedge"}


def _h_nary(el) -> str:
    chr_ = "∑"
    sub = sup = ""
    body = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "naryPr":
            for n2, c2 in ((_local(c2.tag), c2) for c2 in c):
                if n2 == "chr":
                    chr_ = _val(c2) or "∑"
        elif name == "sub":
            sub = omath_to_latex(c)
        elif name == "sup":
            sup = omath_to_latex(c)
        elif name == "e":
            body = omath_to_latex(c)
    op = _NARY_OPS.get(chr_) or r"\mathop{%s}" % chr_
    limits = r"\limits" if op in _NARY_LIMITS_OPS else ""
    return ("%s%s_{%s}^{%s} %s" % (op, limits, sub, sup, body)).rstrip()


def _h_rad(el) -> str:
    deg = base = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "deg":
            deg = omath_to_latex(c)
        elif name == "e":
            base = omath_to_latex(c)
    if deg:
        return r"\sqrt[%s]{%s}" % (deg, base)
    return r"\sqrt{%s}" % base


_DELIM_CHARS = {
    "{": r"\{",
    "}": r"\}",
    "|": "|",
    "‖": r"\|",
    "[": "[",
    "]": "]",
    "(": "(",
    ")": ")",
    "⟨": r"\langle",
    "⟩": r"\rangle",
}


def _delim(c: str) -> str:
    if not c:
        return "."
    return _DELIM_CHARS.get(c, c)


def _h_d(el) -> str:
    beg, end = "(", ")"
    inner_parts = []
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "dPr":
            for n2, c2 in ((_local(c2.tag), c2) for c2 in c):
                if n2 == "begChr":
                    beg = _val(c2)
                elif n2 == "endChr":
                    end = _val(c2)
        else:
            inner_parts.append(omath_to_latex(c))
    inner = "".join(inner_parts)
    return r"\left%s %s \right%s" % (_delim(beg), inner, _delim(end))


def _h_limLow(el) -> str:
    base = lim = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "e":
            base = omath_to_latex(c)
        elif name == "lim":
            lim = omath_to_latex(c)
    return r"\underset{%s}{%s}" % (lim, base)


def _h_limUpp(el) -> str:
    base = lim = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "e":
            base = omath_to_latex(c)
        elif name == "lim":
            lim = omath_to_latex(c)
    return r"\overset{%s}{%s}" % (lim, base)


def _h_bar(el) -> str:
    pos = "top"
    base = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "barPr":
            for n2, c2 in ((_local(c2.tag), c2) for c2 in c):
                if n2 == "pos":
                    pos = _val(c2) or "top"
        elif name == "e":
            base = omath_to_latex(c)
    if pos == "bottom":
        return r"\underline{%s}" % base
    return r"\overline{%s}" % base


_ACC_COMMANDS = {
    "\u0307": r"\dot",  # combining dot above
    "\u0308": r"\ddot",
    "\u0302": r"\hat",  # combining circumflex
    "\u0303": r"\tilde",
    "\u0304": r"\bar",  # combining macron
    "\u030c": r"\check",
    "\u0301": r"\acute",
    "\u0300": r"\grave",
    "˙": r"\dot",
    "̈": r"\ddot",
    "̂": r"\hat",
    "̃": r"\tilde",
    "‾": r"\bar",
    "ˇ": r"\check",
    "′": r"^{\prime}",
}


def _h_acc(el) -> str:
    acc_chr = ""
    base = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "accPr":
            for n2, c2 in ((_local(c2.tag), c2) for c2 in c):
                if n2 == "accChr":
                    acc_chr = _val(c2) or ""
        elif name == "e":
            base = omath_to_latex(c)
    cmd = _ACC_COMMANDS.get(acc_chr, r"\hat")
    return "%s{%s}" % (cmd, base)


def _h_func(el) -> str:
    fname = arg = ""
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "fName":
            fname = omath_to_latex(c)
        elif name == "e":
            arg = omath_to_latex(c)
    return r"%s\left(%s\right)" % (fname, arg)


def _h_m(el) -> str:
    rows = []
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "mr":
            cells = [omath_to_latex(c2) for c2 in c if _local(c2.tag) == "e"]
            rows.append(" & ".join(cells))
    return r"\begin{pmatrix} %s \end{pmatrix}" % r" \\ ".join(rows)


def _h_eqArr(el) -> str:
    rows = [omath_to_latex(c) for c in el if _local(c.tag) == "e"]
    if len(rows) <= 1:
        return rows[0] if rows else ""
    return r"\begin{aligned} %s \end{aligned}" % r" \\ ".join(rows)


def _h_groupChr(el) -> str:
    beg, end = "", ""
    inner_parts = []
    for name, c in ((_local(c.tag), c) for c in el):
        if name == "groupChrPr":
            for n2, c2 in ((_local(c2.tag), c2) for c2 in c):
                if n2 == "begChr":
                    beg = _val(c2)
                elif n2 == "endChr":
                    end = _val(c2)
        else:
            inner_parts.append(omath_to_latex(c))
    inner = "".join(inner_parts)
    return r"\left%s %s \right%s" % (_delim(beg or "("), inner, _delim(end or ")"))


def _h_box(el) -> str:
    return omath_to_latex(el)


def _h_ph(el) -> str:
    return omath_to_latex(el)


def _h_borderBox(el) -> str:
    return omath_to_latex(el)


def _h_none(el) -> str:
    return ""


_OMATH_HANDLERS = {
    "t": _h_t,
    "r": _h_r,
    "f": _h_f,
    "sSup": _h_sSup,
    "sSub": _h_sSub,
    "sSubSup": _h_sSubSup,
    "sPre": _h_sPre,
    "nary": _h_nary,
    "rad": _h_rad,
    "d": _h_d,
    "limLow": _h_limLow,
    "limUpp": _h_limUpp,
    "bar": _h_bar,
    "acc": _h_acc,
    "func": _h_func,
    "m": _h_m,
    "eqArr": _h_eqArr,
    "groupChr": _h_groupChr,
    "box": _h_box,
    "ph": _h_ph,
    "borderBox": _h_borderBox,
    # properties and containers we never render
    "rPr": _h_none, "naryPr": _h_none, "fPr": _h_none, "dPr": _h_none,
    "sSupPr": _h_none, "sSubPr": _h_none, "sSubSupPr": _h_none,
    "radPr": _h_none, "barPr": _h_none, "accPr": _h_none, "funcPr": _h_none,
    "limLowPr": _h_none, "limUppPr": _h_none, "eqArrPr": _h_none, "mPr": _h_none,
    "ctrlPr": _h_none, "argPr": _h_none, "groupChrPr": _h_none, "boxPr": _h_none,
    "borderBoxPr": _h_none, "mathPr": _h_none, "oMathParaPr": _h_none,
    "deg": _h_none, "num": _h_none, "den": _h_none, "e": _h_none, "fName": _h_none,
}


# --- image formulas (WMF/EMF -> PNG -> pix2text) --------------------------

def _convert_to_png(blob: bytes, ext: str) -> bytes | None:
    """Rasterize WMF/EMF to PNG via headless LibreOffice. Returns PNG bytes."""
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / ("formula." + ext)
        out = Path(tmp) / "out"
        out.mkdir()
        src.write_bytes(blob)
        try:
            subprocess.run(
                ["soffice", "-env:UserInstallation=file://%s/lo" % tmp, "--headless",
                 "--convert-to", "png", "--outdir", str(out), str(src)],
                capture_output=True, timeout=90,
            )
        except Exception as e:
            logger.warning("LibreOffice WMF->PNG failed: %s", e)
            return None
        png = out / "formula.png"
        if png.exists():
            return png.read_bytes()
        logger.warning("LibreOffice produced no PNG for %s image", ext)
        return None


def _ocr_formula_image(png: bytes) -> str | None:
    """Recognize a single formula image. Returns LaTeX or None."""
    try:
        import formula_ocr
        latex = formula_ocr.recognize_formula_image(png)
        return latex.strip() or None
    except Exception as e:
        logger.warning("pix2text formula recognition failed: %s", e)
        return None


def _image_latex(rid: str, rels: dict, media: dict, ocr_formulas: bool) -> str | None:
    """Try to OCR the media image referenced by rid. Returns LaTeX or None."""
    target = rels.get(rid)
    if not target:
        return None
    ext = target.rsplit(".", 1)[-1].lower()
    if ext not in _FORMULA_IMAGE_EXTS or not ocr_formulas:
        return None
    blob = media.get("word/" + target) or media.get(target)
    if not blob:
        return None
    png = _convert_to_png(blob, ext)
    if not png:
        return None
    return _ocr_formula_image(png)


# --- DOCX paragraph/table walk --------------------------------------------

def _run_text(r, rels: dict, media: dict, ocr_formulas: bool) -> str:
    out = []
    for c in r:
        name = _local(c.tag)
        if name == "t":
            out.append(c.text or "")
        elif name == "tab":
            out.append("\t")
        elif name == "br":
            out.append("\n")
        elif name in ("drawing", "pict", "object"):
            for a in c.iter():
                lname = _local(a.tag)
                if lname == "blip":
                    rid = a.get("{%s}embed" % R_NS) or a.get("{%s}link" % R_NS)
                elif lname == "imagedata":
                    rid = a.get("{%s}id" % R_NS)
                else:
                    continue
                if not rid:
                    continue
                latex = _image_latex(rid, rels, media, ocr_formulas)
                if latex:
                    out.append("$$" + latex + "$$")
                break  # first usable image per drawing
    return "".join(out)


def _node_text(el, rels: dict, media: dict, ocr_formulas: bool) -> str:
    name = _local(el.tag)
    if name == "r":
        return _run_text(el, rels, media, ocr_formulas)
    if name in ("oMath", "oMathPara"):
        latex = omath_to_latex(el).strip()
        return "$$%s$$" % latex if latex else ""
    if name in ("hyperlink", "ins", "del"):
        return "".join(_node_text(c, rels, media, ocr_formulas) for c in el)
    if name in ("tbl",):
        return _table_text(el, rels, media, ocr_formulas)
    return ""


def _para_text(p, rels: dict, media: dict, ocr_formulas: bool) -> str:
    return "".join(_node_text(c, rels, media, ocr_formulas) for c in p).strip()


def _table_text(tbl, rels: dict, media: dict, ocr_formulas: bool) -> str:
    rows = []
    for tr in tbl:
        if _local(tr.tag) != "tr":
            continue
        cells = []
        for tc in tr:
            if _local(tc.tag) != "tc":
                continue
            cell_parts = [
                _para_text(p, rels, media, ocr_formulas)
                for p in tc
                if _local(p.tag) == "p"
            ]
            cells.append("\n".join(x for x in cell_parts if x))
        rows.append(" | ".join(cells))
    return "\n".join(r for r in rows if r)


def _parse_rels(rels_xml: bytes) -> dict:
    rels = {}
    try:
        import lxml.etree as ET
        root = ET.fromstring(rels_xml)
        for rel in root:
            rels[rel.get("Id")] = rel.get("Target") or ""
    except Exception:
        pass
    return rels


def extract_docx_text(path: Path, ocr_formulas: bool = True) -> str:
    """Full text of a DOCX with inline LaTeX formulas (oMath + WMF/EMF OCR).

    Falls back to plain python-docx paragraph text on parse errors.
    """
    try:
        import lxml.etree as ET
        with zipfile.ZipFile(path) as z:
            doc_xml = z.read("word/document.xml")
            try:
                rels_xml = z.read("word/_rels/document.xml.rels")
            except KeyError:
                rels_xml = b"<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'/>"
            media = {n: z.read(n) for n in z.namelist() if n.startswith("word/media/")}
    except Exception as e:
        logger.warning("Cannot open docx as zip (%s); falling back to python-docx", e)
        return _docx_plain(path)

    try:
        root = ET.fromstring(doc_xml)
    except Exception as e:
        logger.warning("Cannot parse document.xml (%s); falling back to python-docx", e)
        return _docx_plain(path)

    rels = _parse_rels(rels_xml)
    parts = []
    body = None
    for c in root:
        if _local(c.tag) == "body":
            body = c
            break
    if body is None:
        return _docx_plain(path)
    for child in body:
        name = _local(child.tag)
        if name == "p":
            text = _para_text(child, rels, media, ocr_formulas)
        elif name == "tbl":
            text = _table_text(child, rels, media, ocr_formulas)
        else:
            continue
        if text:
            parts.append(text)
    return "\n".join(parts)


def _docx_plain(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        return path.read_text(encoding="utf-8", errors="replace")


# --- RTF ------------------------------------------------------------------

_PICT_WMF_RE = re.compile(
    r"\{\\pict[^{}]*\\wmetafile8[^{}]*?\s+([0-9a-fA-F ]+)\}", re.S
)
_TOKEN = "\x00F{}"


def _rtf_wmf_latex(hex_str: str) -> str | None:
    hex_clean = re.sub(r"\s+", "", hex_str)
    if len(hex_clean) < 4 or len(hex_clean) % 2 != 0:
        return None
    try:
        blob = bytes.fromhex(hex_clean)
    except ValueError:
        return None
    png = _convert_to_png(blob, "wmf")
    if not png:
        return None
    return _ocr_formula_image(png)


def extract_rtf_with_math(path: Path, ocr_formulas: bool = True) -> str:
    """RTF text with `{\\pict\\wmetafile8 ...}` formulas OCR'd in place."""
    try:
        raw = path.read_bytes().decode("latin-1")
    except Exception:
        return ""
    formulas: list[str | None] = []

    def _repl(m: re.Match) -> str:
        if ocr_formulas:
            formulas.append(_rtf_wmf_latex(m.group(1)))
        else:
            formulas.append(None)
        return _TOKEN.format(len(formulas))

    processed = _PICT_WMF_RE.sub(_repl, raw)
    text = _strip_rtf(processed)
    for i, latex in enumerate(formulas):
        token = _TOKEN.format(i + 1)
        if latex:
            text = text.replace(token, "$$%s$$" % latex)
        else:
            text = text.replace(token, "")
    return text


def _strip_rtf(raw: str) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
        return rtf_to_text(raw)
    except Exception:
        pass
    text = re.sub(r"\\'([0-9a-fA-F]{2})", lambda m: _hex_to_char(m.group(1)), raw)
    text = re.sub(r"\\[a-zA-Z]+\d*\s?", " ", text)
    text = re.sub(r"\{[^{}]*\}", "", text)
    text = text.replace("\\par", "\n").replace("\\cell", " | ").replace("\\row", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text


def _hex_to_char(hex_str: str) -> str:
    try:
        return bytes([int(hex_str, 16)]).decode("cp1251")
    except Exception:
        return "?"
